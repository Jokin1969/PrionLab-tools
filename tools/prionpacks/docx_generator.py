import io
import logging
import re
from base64 import b64decode
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import members as members_module

logger = logging.getLogger(__name__)

_TEAL      = RGBColor(0x1a, 0x73, 0x73)
_DARK      = RGBColor(0x1e, 0x2d, 0x3d)
_DIM       = RGBColor(0x64, 0x74, 0x8b)
_LIGHT_BG  = RGBColor(0xf0, 0xf7, 0xf7)
_WHITE     = RGBColor(0xff, 0xff, 0xff)

# Accent colors per responsible — mirrored from CSS .pp-pkg-card-id rules
_RESPONSIBLE_ACCENT: dict[str, tuple[str, RGBColor]] = {
    'joaquin': ('1d4ed8', RGBColor(0x1d, 0x4e, 0xd8)),
    'hasier':  ('15803d', RGBColor(0x15, 0x80, 0x3d)),
    'jorge':   ('c2410c', RGBColor(0xc2, 0x41, 0x0c)),
    'carlos':  ('7c3aed', RGBColor(0x7c, 0x3a, 0xed)),
}
_TEAL_HEX = '1a7373'


def _accent(resp_id: str) -> tuple[str, RGBColor]:
    """Return (hex_str, RGBColor) for the responsible person's accent colour."""
    return _RESPONSIBLE_ACCENT.get((resp_id or '').lower(), (_TEAL_HEX, _TEAL))

# Markdown-style superscript: PrP^Sc^ -> PrP + superscript("Sc")
_SUP_RE = re.compile(r'\^([^\^\s][^\^]*?)\^')

# Header line that introduces the abstract block inside a reference string
_REF_SUMMARY_HEADERS = re.compile(
    r'^\s*(Resumen|Resumen del artículo|Abstract)\s*:?\s*$',
    re.IGNORECASE | re.MULTILINE,
)

# DOI pattern: matches "DOI: 10.xxxx/..." (case-insensitive)
_DOI_RE = re.compile(r'(DOI:\s*)(10\.\d{4,}[./][^\s,;]+)', re.IGNORECASE)

# Inline DOI pattern: bare DOIs and doi.org URLs in prose text
_INLINE_DOI_RE = re.compile(
    r'(?:https?://(?:dx\.)?doi\.org/|doi\.org/)?'
    r'\b(10\.\d{4,}/[^\s,;>\]\)]+)',
    re.IGNORECASE,
)


def _apply_font(run, *, size=None, bold=None, italic=None, color=None):
    if size is not None:   run.font.size = size
    if bold is not None:   run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None:  run.font.color.rgb = color


def add_runs(paragraph, text, *, size=None, bold=None, italic=None, color=None):
    """add_run() replacement that turns ^xxx^ markers into superscript runs.

    All runs share the supplied font attributes; only `superscript` differs
    on the bracketed segments.
    """
    if not text:
        return
    pos = 0
    for m in _SUP_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            _apply_font(r, size=size, bold=bold, italic=italic, color=color)
        rs = paragraph.add_run(m.group(1))
        rs.font.superscript = True
        _apply_font(rs, size=size, bold=bold, italic=italic, color=color)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        _apply_font(r, size=size, bold=bold, italic=italic, color=color)

def _hyperlink_url(para, display: str, url: str, *, pt: int = 10):
    """Add an external URL hyperlink run to `para` (blue underlined)."""
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u');     u.set(qn('w:val'), 'single')
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(pt * 2))
    rpr.append(c); rpr.append(u); rpr.append(sz)
    t = OxmlElement('w:t'); t.text = display
    r.append(rpr); r.append(t)
    hl.append(r)
    para._p.append(hl)


def add_runs_with_doi(paragraph, text, *, size=None, bold=None, italic=None, color=None):
    """Like add_runs() but turns DOI values into clickable doi.org hyperlinks."""
    if not text:
        return
    pt = int(size.pt) if size else 10
    last = 0
    for m in _DOI_RE.finditer(text):
        if m.start() > last:
            add_runs(paragraph, text[last:m.start()], size=size, bold=bold, italic=italic, color=color)
        label_run = paragraph.add_run(m.group(1))
        _apply_font(label_run, size=size, bold=bold, italic=italic, color=color)
        doi_val = m.group(2).rstrip('.,;)')
        _hyperlink_url(paragraph, doi_val, f'https://doi.org/{doi_val}', pt=pt)
        last = m.end()
    if last < len(text):
        add_runs(paragraph, text[last:], size=size, bold=bold, italic=italic, color=color)


def add_runs_with_inline_doi(paragraph, text, *, size=None, bold=None, italic=None, color=None):
    """Like add_runs() but turns bare DOIs and doi.org URLs in prose into hyperlinks."""
    if not text:
        return
    pt = int(size.pt) if size else 10
    last = 0
    for m in _INLINE_DOI_RE.finditer(text):
        if m.start() > last:
            add_runs(paragraph, text[last:m.start()], size=size, bold=bold, italic=italic, color=color)
        doi_val = m.group(1).rstrip('.,;)')
        _hyperlink_url(paragraph, m.group(0).rstrip('.,;)'), f'https://doi.org/{doi_val}', pt=pt)
        last = m.end()
    if last < len(text):
        add_runs(paragraph, text[last:], size=size, bold=bold, italic=italic, color=color)


PRIORITY_ES = {'high': 'Alta', 'medium': 'Media', 'low': 'Baja', 'none': '—'}
TYPE_ES = {
    'research': 'Investigación', 'review': 'Revisión',
    'clinical': 'Ensayo clínico', 'case': 'Caso clínico', 'meta': 'Meta-análisis',
}


def _split_reference(ref: str) -> tuple:
    """Return (header, abstract) by splitting at a 'Resumen:'/'Abstract:' line.

    If no such marker exists the full text is returned as header with an
    empty abstract string.
    """
    if not ref:
        return '', ''
    m = _REF_SUMMARY_HEADERS.search(ref)
    if not m:
        return ref.rstrip(), ''
    return ref[:m.start()].rstrip(), ref[m.end():].strip()


def _set_compat_mode_15(doc: Document):
    """Upgrade document to Word 2013 compat mode and fix app metadata.

    python-docx's default template ships with AppVersion=14 (Word 2010 for Mac).
    Word 365 sees that and applies Word 2010 behaviour, silently ignoring
    w:collapsed.  We patch both settings.xml (compatibilityMode=15) and
    docProps/app.xml (AppVersion=16, Application=Microsoft Office Word) so
    Word 365 trusts and honours the collapse state.
    """
    # --- settings.xml: compatibilityMode=15 ---
    settings_el = doc.settings.element
    compat = settings_el.find(qn('w:compat'))
    if compat is not None:
        for cs in compat.findall(qn('w:compatSetting')):
            if cs.get(qn('w:name')) == 'compatibilityMode':
                cs.set(qn('w:val'), '15')
                break
        else:
            cs = OxmlElement('w:compatSetting')
            cs.set(qn('w:name'), 'compatibilityMode')
            cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
            cs.set(qn('w:val'), '15')
            compat.append(cs)
    else:
        compat = OxmlElement('w:compat')
        cs = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'), '15')
        compat.append(cs)
        settings_el.append(compat)



def _patch_docx(docx_bytes: bytes) -> bytes:
    """Post-process the raw .docx bytes (ZIP) to fix two things:

    1. docProps/app.xml — claim Word 2016 (AppVersion=16.0000) so Word 365
       doesn't apply legacy Word-2010-for-Mac behaviour.

    2. word/document.xml — replace every ``<w:collapsed w:val="1"/>`` with
       ``<w15:collapsed w:val="1"/>`` (the Word-2013 extension namespace).
       The standard w:collapsed is often silently ignored by Word desktop;
       w15:collapsed is reliably respected.

    All other ZIP members are copied verbatim so images, relationships,
    styles, numbering, headers/footers etc. are untouched.
    """
    from lxml import etree as _et
    import zipfile as _zf

    _W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _W15_NS = 'http://schemas.microsoft.com/office/word/2012/wordml'
    _MC_NS  = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    _APP_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'

    def _patch_app(data: bytes) -> bytes:
        root = _et.fromstring(data)
        for tag, val in (('Application', 'Microsoft Office Word'),
                         ('AppVersion', '16.0000')):
            el = root.find(f'{{{_APP_NS}}}{tag}')
            if el is not None:
                el.text = val
            else:
                _et.SubElement(root, f'{{{_APP_NS}}}{tag}').text = val
        return _et.tostring(root, xml_declaration=True,
                            encoding='UTF-8', standalone=True)

    def _patch_document(data: bytes) -> bytes:
        root = _et.fromstring(data)

        # Ensure w15 namespace is declared on the root and listed in mc:Ignorable.
        ignorable = root.get(f'{{{_MC_NS}}}Ignorable', '')
        if 'w15' not in ignorable.split():
            root.set(f'{{{_MC_NS}}}Ignorable', (ignorable + ' w15').strip())

        # Replace every <w:collapsed w:val="1"/> with <w15:collapsed w:val="1"/>.
        for old in root.findall(f'.//{{{_W_NS}}}collapsed'):
            parent = old.getparent()
            idx    = list(parent).index(old)
            parent.remove(old)
            new_el = _et.Element(f'{{{_W15_NS}}}collapsed')
            new_el.set(f'{{{_W_NS}}}val', '1')
            parent.insert(idx, new_el)

        return _et.tostring(root, xml_declaration=True,
                            encoding='UTF-8', standalone=True)

    try:
        in_buf  = io.BytesIO(docx_bytes)
        out_buf = io.BytesIO()
        with _zf.ZipFile(in_buf, 'r') as zin, \
             _zf.ZipFile(out_buf, 'w', compression=_zf.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'docProps/app.xml':
                    data = _patch_app(data)
                elif item.filename == 'word/document.xml':
                    data = _patch_document(data)
                zout.writestr(item, data)
        return out_buf.getvalue()
    except Exception:
        return docx_bytes  # best-effort; return original on any error


def generate_package_docx(pkg: dict, version: int, send_date: datetime) -> bytes:
    doc = Document()

    # Upgrade compatibility to Word 2013+ so w:collapsed on headings is honoured.
    _set_compat_mode_15(doc)

    # ── Page margins ────────────────────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # ── Resolve responsible + accent colour ────────────────────────────────────────────────
    resp_id = pkg.get('responsible') or ''
    resp_name = ''
    if resp_id:
        m = members_module.get_member(resp_id)
        if m:
            resp_name = f"{m['name']} {m['surname']}"
    acc_hex, ACCENT = _accent(resp_id)

    # ── Page header: PRP-### · Responsable · date ───────────────────────────────────────────
    hdr = sec.header
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_id = hp.add_run(pkg.get('id', 'PRP'))
    r_id.font.bold = True; r_id.font.size = Pt(8); r_id.font.color.rgb = ACCENT
    if resp_name:
        r_resp = hp.add_run(f'  ·  {resp_name}')
        r_resp.font.size = Pt(8); r_resp.font.color.rgb = _DIM
    r_sep = hp.add_run('  ·  ' + send_date.strftime('%d/%m/%Y'))
    r_sep.font.size = Pt(8); r_sep.font.color.rgb = _DIM

    # ── Title ────────────────────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(t, pkg.get('title', 'Sin título'), size=Pt(14), bold=True, color=ACCENT)

    # Subtitle must come BEFORE alt titles so it is never inside the collapsed section.
    resp_part = f"  ·  Responsable: {resp_name}" if resp_name else ''
    sub = doc.add_paragraph()
    run2 = sub.add_run(
        f"PrionPack {pkg.get('id', '')}  ·  "
        f"Versión {version}{resp_part}  ·  "
        f"Generado el {send_date.strftime('%d/%m/%Y %H:%M')}"
    )
    run2.font.size  = Pt(9)
    run2.font.color.rgb = _DIM

    doc.add_paragraph()

    # Alternative titles — collapsed section so they don't clutter the opening view
    alt_titles = [at.strip() for at in (pkg.get('altTitles') or []) if (at or '').strip()]
    if alt_titles:
        _section_heading(doc, 'TÍTULOS ALTERNATIVOS', collapsed=True, accent=ACCENT, accent_hex=acc_hex)
        for at_clean in alt_titles:
            # Use Heading 3 so these paragraphs have outlineLvl=2 and are
            # definitively inside the Heading 2 collapse scope.  Override
            # spacing/formatting so they look like plain body text.
            ap = doc.add_paragraph(style='Heading 3')
            ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after  = Pt(2)
            ap.paragraph_format.keep_with_next = False
            add_runs(ap, at_clean, size=Pt(12), italic=True, bold=False, color=ACCENT)
        doc.add_paragraph(style='Heading 3').paragraph_format.space_before = Pt(0)

    def sh(text, collapsed=True, **kw):
        _section_heading(doc, text, collapsed=collapsed, accent=ACCENT, accent_hex=acc_hex, **kw)

    # ── Description ──────────────────────────────────────────────────────────────────────────
    desc = (pkg.get('description') or '').strip()
    if desc:
        sh('DESCRIPCIÓN BREVE')
        p = doc.add_paragraph()
        add_runs(p, desc, italic=True, size=Pt(10), color=_DIM)
        doc.add_paragraph()

    # ── Co-authors ──────────────────────────────────────────────────────────────────────────
    coauthors = (pkg.get('coAuthors') or '').strip()
    if coauthors:
        sh('CO-AUTORES')
        p = doc.add_paragraph()
        add_runs(p, coauthors, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Affiliations ─────────────────────────────────────────────────────────────────────────
    affiliations = (pkg.get('affiliations') or '').strip()
    if affiliations:
        sh('AFILIACIONES')
        p = doc.add_paragraph()
        add_runs(p, affiliations, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Abstract ──────────────────────────────────────────────────────────────────────────────
    abstract = (pkg.get('abstract') or '').strip()
    if abstract:
        sh('ABSTRACT', collapsed=True)
        p = doc.add_paragraph()
        add_runs(p, abstract, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Author Summary ─────────────────────────────────────────────────────────────────────────
    author_summary = (pkg.get('authorSummary') or '').strip()
    if author_summary:
        sh('RESUMEN PARA AUTORES', collapsed=True)
        p = doc.add_paragraph()
        add_runs(p, author_summary, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Introduction ─────────────────────────────────────────────────────────────────────────
    intro = (pkg.get('introduction') or '').strip()
    if intro:
        sh('INTRODUCCIÓN')
        p = doc.add_paragraph()
        add_runs_with_inline_doi(p, intro, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Introduction References (Ri-XX) ───────────────────────────────────────────────────────
    intro_refs_raw = pkg.get('introReferences')
    if isinstance(intro_refs_raw, list):
        intro_refs_list = [str(r).strip() for r in intro_refs_raw if isinstance(r, str) and r.strip()]
    elif isinstance(intro_refs_raw, str) and intro_refs_raw.strip():
        intro_refs_list = [intro_refs_raw.strip()]
    else:
        intro_refs_list = []
    if intro_refs_list:
        sh('REFERENCIAS DE INTRODUCCIÓN', collapsed=True)
        for i, ref in enumerate(intro_refs_list, 1):
            header, abstract = _split_reference(ref)
            _make_ref_heading(doc, f'[Ri-{i:02d}] ', header, ACCENT, add_runs_with_doi)
            if abstract:
                _ref_abstract_para(doc, abstract, _DIM)

    # ── Methods (multi-field: list of {title, body}) ──────────────────────────────────────────
    methods_raw = pkg.get('methods')
    if isinstance(methods_raw, list):
        methods_list = []
        for m in methods_raw:
            if isinstance(m, dict):
                t = (m.get('title') or '').strip()
                b = (m.get('body') or '').strip()
                if t or b:
                    methods_list.append({'title': t, 'body': b})
            elif isinstance(m, str) and m.strip():
                methods_list.append({'title': '', 'body': m.strip()})
    elif isinstance(methods_raw, str) and methods_raw.strip():
        methods_list = [{'title': '', 'body': methods_raw.strip()}]
    else:
        methods_list = []
    if methods_list:
        sh('MÉTODOS')
        for mi, m in enumerate(methods_list, 1):
            if m['title']:
                p_t = doc.add_paragraph()
                r_n = p_t.add_run(f'M-{mi:02d} — ')
                r_n.font.bold = True; r_n.font.size = Pt(11); r_n.font.color.rgb = ACCENT
                add_runs(p_t, m['title'], size=Pt(11), bold=True, color=ACCENT)
            else:
                p_t = doc.add_paragraph()
                r_n = p_t.add_run(f'M-{mi:02d}')
                r_n.font.bold = True; r_n.font.size = Pt(11); r_n.font.color.rgb = ACCENT
            if m['body']:
                p_b = doc.add_paragraph()
                add_runs(p_b, m['body'], size=Pt(10), color=_DARK)
            doc.add_paragraph()

    # ── Investigations ──────────────────────────────────────────────────────────────────────────
    inv = pkg.get('investigations') or {}
    inv_text  = (inv.get('text') or '').strip()
    inv_files = inv.get('files') or []
    if inv_text or inv_files:
        sh('INVESTIGACIONES')
        if inv_text:
            p = doc.add_paragraph()
            add_runs(p, inv_text, size=Pt(10), color=_DARK)
        if inv_files:
            p = doc.add_paragraph()
            r = p.add_run('Documentos adjuntos:')
            r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = _DARK
            for f in inv_files:
                p2 = doc.add_paragraph(style='List Bullet')
                add_runs(p2, f.get('name', 'documento'), size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Pre-process gaps for linking ────────────────────────────────────────────────
    raw_missing = pkg.get('gaps', {}).get('missingInfo', [])
    gap_items = []
    for i, g in enumerate(raw_missing):
        if isinstance(g, str):
            gap_items.append({'text': g, 'fid': None, 'bm': f'ppgap_{i}', 'neededExperiment': ''})
        else:
            gap_items.append({
                'text': g.get('text', ''),
                'fid': g.get('findingId'),
                'bm': f'ppgap_{i}',
                'neededExperiment': g.get('neededExperiment') or '',
            })

    gaps_for_finding: dict = {}
    for gi in gap_items:
        if gi['fid']:
            gaps_for_finding.setdefault(gi['fid'], []).append(gi)

    # ── Findings ─────────────────────────────────────────────────────────────────────────────
    findings = pkg.get('findings', [])
    if findings:
        sh('HALLAZGOS PRINCIPALES')
        for fi, finding in enumerate(findings, 1):
            fid = finding.get('id', '')
            p = doc.add_paragraph()
            bm_name = f'ppfinding_{fid}' if fid else f'ppfinding_{fi}'
            _bookmark_add(p, bm_name, fi * 100)
            r_prefix = p.add_run(f'▶  F-{fi:02d} — ')
            r_prefix.font.bold = True; r_prefix.font.size = Pt(11); r_prefix.font.color.rgb = ACCENT
            add_runs(p, finding.get("title", ""), size=Pt(11), bold=True, color=ACCENT)

            en_title = (finding.get('titleEnglish') or '').strip()
            if en_title:
                p2 = doc.add_paragraph()
                add_runs(p2, en_title, italic=True, size=Pt(9), color=_DIM)

            fdesc = (finding.get('description') or '').strip()
            if fdesc:
                p3 = doc.add_paragraph()
                add_runs(p3, fdesc, size=Pt(10), color=_DARK)

            for figi, fig in enumerate(finding.get('figures', []), 1):
                _render_figure(doc, fig, fi, figi, accent=ACCENT)

            for tabi, tbl_item in enumerate(finding.get('tables', []), 1):
                p = doc.add_paragraph()
                r = p.add_run(f'Tabla {fi}.{tabi}')
                r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = ACCENT
                desc_t = (tbl_item.get('description') or '').strip()
                if desc_t:
                    r_sep = p.add_run('  —  ')
                    r_sep.font.size = Pt(9); r_sep.font.color.rgb = _DIM
                    add_runs(p, desc_t, size=Pt(9), color=_DIM)

            linked = gaps_for_finding.get(fid, [])
            if linked:
                p_gaps = doc.add_paragraph()
                r_label = p_gaps.add_run('Gaps asociados: ')
                r_label.font.size = Pt(9); r_label.font.bold = True; r_label.font.color.rgb = _DIM
                for k, gi in enumerate(linked):
                    if k: p_gaps.add_run(', ').font.size = Pt(9)
                    _hyperlink_anchor(p_gaps, gi['text'], gi['bm'], pt=9)

            doc.add_paragraph()

    # ── Gaps & Next Steps ─────────────────────────────────────────────────────────────────
    sh('GAPS & NEXT STEPS')

    if gap_items:
        p_h = doc.add_paragraph()
        r_h = p_h.add_run('Información faltante')
        r_h.font.bold = True; r_h.font.size = Pt(10); r_h.font.color.rgb = _DARK

        for gi in gap_items:
            p = doc.add_paragraph()
            _bookmark_add(p, gi['bm'], abs(hash(gi['bm'])) % 90000 + 1000)
            r_miss = p.add_run('Missing: ')
            r_miss.font.bold = True; r_miss.font.size = Pt(10); r_miss.font.color.rgb = _DARK
            add_runs(p, gi['text'], size=Pt(10), color=_DARK)

            needed_exp = (gi.get('neededExperiment') or '').strip()
            if needed_exp:
                p_ne = doc.add_paragraph()
                r_ne = p_ne.add_run('     → Needed: ')
                r_ne.font.size = Pt(9); r_ne.font.italic = True; r_ne.font.color.rgb = _DIM
                add_runs(p_ne, needed_exp, size=Pt(9), italic=True, color=_DIM)

            if gi['fid']:
                linked_f = next((f for f in findings if f.get('id') == gi['fid']), None)
                if linked_f:
                    fi_num = findings.index(linked_f) + 1
                    ann = doc.add_paragraph()
                    r_ann = ann.add_run(f'     → Vinculado a F-{fi_num:02d}: ')
                    r_ann.font.size = Pt(8); r_ann.font.italic = True; r_ann.font.color.rgb = ACCENT
                    add_runs(ann, linked_f.get("title", ""), size=Pt(8), italic=True, color=ACCENT)

    doc.add_paragraph()

    # ── Discussion ──────────────────────────────────────────────────────────────────────────
    disc = (pkg.get('discussion') or '').strip()
    if disc:
        sh('DISCUSIÓN')
        p = doc.add_paragraph()
        add_runs_with_inline_doi(p, disc, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Acknowledgments ───────────────────────────────────────────────────────────────────────
    acknowledgments = (pkg.get('acknowledgments') or '').strip()
    if acknowledgments:
        sh('AGRADECIMIENTOS')
        p = doc.add_paragraph()
        add_runs(p, acknowledgments, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Funding ────────────────────────────────────────────────────────────────────────────
    funding = (pkg.get('funding') or '').strip()
    if funding:
        sh('FINANCIACIÓN')
        p = doc.add_paragraph()
        add_runs(p, funding, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Conflicts of Interest ───────────────────────────────────────────────────────────────
    conflicts = (pkg.get('conflictsOfInterest') or '').strip()
    if conflicts:
        sh('CONFLICTOS DE INTERÉS')
        p = doc.add_paragraph()
        add_runs(p, conflicts, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── References ───────────────────────────────────────────────────────────────────────────
    refs_raw = pkg.get('references')
    if isinstance(refs_raw, list):
        refs_list = [str(r).strip() for r in refs_raw if isinstance(r, str) and r.strip()]
    elif isinstance(refs_raw, str) and refs_raw.strip():
        refs_list = [refs_raw.strip()]
    else:
        refs_list = []
    if refs_list:
        sh('REFERENCIAS', collapsed=True)
        for i, ref in enumerate(refs_list, 1):
            header, abstract = _split_reference(ref)
            _make_ref_heading(doc, f'[{i}] ', header, ACCENT, add_runs_with_doi)
            if abstract:
                _ref_abstract_para(doc, abstract, _DIM)

    # ── CReDiT ────────────────────────────────────────────────────────────────────────────────
    credit = (pkg.get('credit') or '').strip()
    if credit:
        sh('CONTRIBUCIÓN DE AUTORÍA (CReDiT)')
        p = doc.add_paragraph()
        add_runs(p, credit, size=Pt(10), color=_DARK)
        doc.add_paragraph()

    # ── Footer ──────────────────────────────────────────────────────────────────────────────
    p_ft = doc.add_paragraph()
    r_ft = p_ft.add_run(
        f'Documento generado automáticamente por PrionLab Tools · '
        f'v{version} · {send_date.strftime("%d %b %Y")}'
    )
    r_ft.font.size = Pt(8); r_ft.font.color.rgb = _DIM

    buf = io.BytesIO()
    doc.save(buf)
    return _patch_docx(buf.getvalue())


def generate_packages_list_docx(packages: list, gen_date: datetime) -> bytes:
    """Generate a compact Word catalogue listing all PrionPacks (id + title)."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # Header
    hdr = sec.header
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_h = hp.add_run('PrionPacks · Lista general')
    r_h.font.bold = True; r_h.font.size = Pt(8); r_h.font.color.rgb = _TEAL
    r_hd = hp.add_run('  ·  ' + gen_date.strftime('%d/%m/%Y'))
    r_hd.font.size = Pt(8); r_hd.font.color.rgb = _DIM

    # Document title
    t = doc.add_paragraph()
    add_runs(t, 'Lista de PrionPacks', size=Pt(18), bold=True, color=_TEAL)
    sub = doc.add_paragraph()
    r_sub = sub.add_run(f'Generado el {gen_date.strftime("%d/%m/%Y %H:%M")}  ·  {len(packages)} paquetes')
    r_sub.font.size = Pt(9); r_sub.font.color.rgb = _DIM
    doc.add_paragraph()

    _section_heading(doc, 'PAQUETES DE INFORMACIÓN')

    for pkg in sorted(packages, key=lambda p: p.get('id', '')):
        p = doc.add_paragraph()
        r_id = p.add_run(pkg.get('id', '—') + '  ')
        r_id.font.bold = True; r_id.font.size = Pt(10); r_id.font.color.rgb = _TEAL
        add_runs(p, pkg.get('title', 'Sin título'), size=Pt(10), color=_DARK)
        # Optional: one-line description in dim italic
        desc = (pkg.get('description') or '').strip()
        if desc:
            short_desc = desc[:120] + ('…' if len(desc) > 120 else '')
            p_d = doc.add_paragraph()
            add_runs(p_d, short_desc, size=Pt(9), italic=True, color=_DIM)
            p_d.paragraph_format.left_indent = Cm(0.8)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Footer
    p_ft = doc.add_paragraph()
    r_ft = p_ft.add_run(
        f'Lista generada automáticamente por PrionLab Tools · {gen_date.strftime("%d %b %Y")}'
    )
    r_ft.font.size = Pt(8); r_ft.font.color.rgb = _DIM

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _section_heading(doc: Document, text: str, collapsed: bool = False,
                     accent: RGBColor = None, accent_hex: str = None):
    p = doc.add_paragraph(style='Heading 2')
    if collapsed:
        pPr = p._p.get_or_add_pPr()
        collapsed_el = OxmlElement('w:collapsed')
        collapsed_el.set(qn('w:val'), '1')
        pPr.append(collapsed_el)
    run = p.add_run(text)
    run.font.bold      = True
    run.font.size      = Pt(10)
    run.font.color.rgb = accent if accent is not None else _TEAL
    _para_border_bottom(p, accent_hex or _TEAL_HEX)


def _make_ref_heading(doc: Document, label: str, body: str,
                      accent: RGBColor, add_runs_fn) -> None:
    """Add a collapsible reference heading (Heading 3) + body paragraph.

    Heading 3 keeps: keepNext=0, keepLines=0, pageBreakBefore=0
    so chains of references never force a new page.

    The body paragraph gets an explicit outlineLvl=3 (outline level 4,
    below Heading 3's level 2) so Word includes it in the heading's
    w15:collapsed scope without turning it into a heading-style paragraph.
    """
    _DARK_LOCAL = RGBColor(0x1e, 0x2d, 0x3d)

    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

    # Override inherited keepNext / keepLines / pageBreakBefore
    pPr = p._p.get_or_add_pPr()
    for tag in ('keepNext', 'keepLines', 'pageBreakBefore'):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:val'), '0')
        pPr.insert(1, el)  # insert after pStyle, before anything else

    # Runs
    r_label = p.add_run(label)
    r_label.font.size = Pt(10); r_label.font.bold = True
    r_label.font.color.rgb = accent
    add_runs_fn(p, body, size=Pt(10), color=_DARK_LOCAL)

    # w:collapsed — will be converted to w15:collapsed by _patch_docx
    collapsed_el = OxmlElement('w:collapsed')
    collapsed_el.set(qn('w:val'), '1')
    pPr.append(collapsed_el)


def _ref_abstract_para(doc: Document, text: str, color: RGBColor) -> None:
    """Plain paragraph with outlineLvl=3 so it falls inside a Heading 3
    collapse scope, without appearing as a heading in the nav pane."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), '3')
    pPr.append(ol)
    add_runs(p, text, size=Pt(9), italic=True, color=color)


def _dim_para(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = _DIM


def _render_figure(doc: Document, fig: dict, fi: int, figi: int, accent: RGBColor = None):
    if accent is None:
        accent = _TEAL
    img_url = fig.get('imageUrl') or fig.get('image') or ''
    caption = (fig.get('caption') or fig.get('description') or '').strip()
    label   = f'Figura {fi}.{figi}'

    if img_url and img_url.startswith('data:'):
        try:
            header, b64data = img_url.split(',', 1)
            mime = header.split(';')[0].split(':')[1]
            if 'svg' not in mime:
                img_bytes = b64decode(b64data)
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(io.BytesIO(img_bytes), width=Cm(12))
        except Exception:
            pass

    p_cap = doc.add_paragraph()
    r_label = p_cap.add_run(label)
    r_label.font.bold = True; r_label.font.size = Pt(9); r_label.font.color.rgb = accent
    if caption:
        r_sep = p_cap.add_run('  ')
        r_sep.font.size = Pt(9); r_sep.font.color.rgb = _DIM
        add_runs(p_cap, caption, size=Pt(9), color=_DIM)


def _bookmark_add(para, bm_name: str, bm_id: int):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'),   str(bm_id))
    start.set(qn('w:name'), bm_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bm_id))
    para._p.append(start)
    para._p.append(end)


def _hyperlink_anchor(para, text: str, anchor: str, pt: int = 10):
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1a7373')
    u = OxmlElement('w:u');         u.set(qn('w:val'), 'single')
    sz = OxmlElement('w:sz');       sz.set(qn('w:val'), str(pt * 2))
    rpr.append(color); rpr.append(u); rpr.append(sz)
    t = OxmlElement('w:t'); t.text = text
    r.append(rpr); r.append(t)
    hl.append(r)
    para._p.append(hl)


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcp.append(shd)


def _para_border_bottom(para, hex_color: str = _TEAL_HEX):
    ppr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _fmt_date(iso: str) -> str:
    if not iso:
        return '—'
    try:
        return datetime.fromisoformat(iso.replace('Z', '+00:00')).strftime('%d/%m/%Y')
    except Exception:
        return iso
