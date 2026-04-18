import csv
import logging
import os
import re
import uuid
from datetime import datetime, timedelta

import pandas as pd
from docx import Document

import config

logger = logging.getLogger(__name__)

TEMPLATES_CSV = os.path.join(config.CSV_DIR, "export_templates.csv")
HISTORY_CSV = os.path.join(config.CSV_DIR, "export_history.csv")

TEMPLATES_COLS = ["template_id", "name", "format", "description", "sections_order", "is_active"]
HISTORY_COLS = [
    "export_id", "user_id", "template_id", "sections_exported",
    "filename", "filepath", "created_at", "download_count", "expires_at",
]

_SEED_TEMPLATES = [
    {
        "template_id": "exp_001", "name": "Basic Word", "format": "docx",
        "description": "Standard Word document format",
        "sections_order": "author_order,methods,funding,acknowledgments,credit,competing_interests",
        "is_active": "true",
    },
    {
        "template_id": "exp_002", "name": "Basic LaTeX", "format": "latex",
        "description": "Standard LaTeX academic format",
        "sections_order": "author_order,methods,funding,acknowledgments,credit,competing_interests",
        "is_active": "true",
    },
    {
        "template_id": "exp_003", "name": "Nature Word", "format": "docx",
        "description": "Nature journal Word format with specific styling",
        "sections_order": "author_order,methods,acknowledgments,funding,credit,competing_interests",
        "is_active": "true",
    },
    {
        "template_id": "exp_004", "name": "Nature LaTeX", "format": "latex",
        "description": "Nature journal LaTeX format",
        "sections_order": "author_order,methods,acknowledgments,funding,credit,competing_interests",
        "is_active": "true",
    },
    {
        "template_id": "exp_005", "name": "Methods Only", "format": "docx",
        "description": "Export only Methods section for supplementary",
        "sections_order": "methods",
        "is_active": "true",
    },
]

_SECTION_TITLES = {
    "author_order": "Author Order & Affiliations",
    "methods": "Materials and Methods",
    "funding": "Funding",
    "acknowledgments": "Acknowledgments",
    "credit": "Author Contributions",
    "competing_interests": "Competing Interests",
}


def _read_templates() -> pd.DataFrame:
    if not os.path.exists(TEMPLATES_CSV):
        return pd.DataFrame(columns=TEMPLATES_COLS)
    return pd.read_csv(TEMPLATES_CSV, dtype=str).fillna("")


def _read_history() -> pd.DataFrame:
    if not os.path.exists(HISTORY_CSV):
        return pd.DataFrame(columns=HISTORY_COLS)
    return pd.read_csv(HISTORY_CSV, dtype=str).fillna("")


def _write_history(df: pd.DataFrame) -> None:
    df.to_csv(HISTORY_CSV, index=False, quoting=csv.QUOTE_ALL)


def _seed_export_templates_if_empty() -> None:
    if os.path.exists(TEMPLATES_CSV):
        try:
            if not pd.read_csv(TEMPLATES_CSV, dtype=str).empty:
                return
        except Exception:
            pass
    pd.DataFrame(_SEED_TEMPLATES).to_csv(TEMPLATES_CSV, index=False, quoting=csv.QUOTE_ALL)
    logger.info("Export templates seeded.")


def bootstrap_export_schema() -> None:
    _seed_export_templates_if_empty()
    if not os.path.exists(HISTORY_CSV):
        pd.DataFrame(columns=HISTORY_COLS).to_csv(HISTORY_CSV, index=False, quoting=csv.QUOTE_ALL)


def get_export_templates() -> list:
    df = _read_templates()
    return df[df["is_active"] == "true"].to_dict(orient="records")


def get_export_template(template_id: str) -> dict | None:
    df = _read_templates()
    row = df[df["template_id"] == template_id]
    return row.iloc[0].to_dict() if not row.empty else None


def get_template_sections_order(template_id: str) -> list:
    tmpl = get_export_template(template_id)
    if not tmpl:
        return []
    return [s.strip() for s in tmpl.get("sections_order", "").split(",") if s.strip()]


def format_section_title(key: str) -> str:
    return _SECTION_TITLES.get(key, key.replace("_", " ").title())


def _escape_latex(text: str) -> str:
    special = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    pattern = re.compile("|".join(re.escape(k) for k in special))
    return pattern.sub(lambda m: special[m.group()], text)


def export_to_word(sections_data: dict, template_id: str = "exp_001") -> str:
    doc = Document()
    title = sections_data.get("title", "Manuscript Draft")
    doc.add_heading(title, 0)

    order = get_template_sections_order(template_id)
    if not order:
        order = [k for k in sections_data if k != "title"]

    for key in order:
        if key not in sections_data:
            continue
        text = sections_data[key]
        if isinstance(text, dict):
            text = text.get("text", "")
        if not text:
            continue
        doc.add_heading(format_section_title(key), level=1)
        doc.add_paragraph(text)

    filename = f"manuscript_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join("/tmp", filename)
    doc.save(filepath)
    return filepath


def export_to_latex(sections_data: dict, template_id: str = "exp_002") -> str:
    title = _escape_latex(sections_data.get("title", "Manuscript Draft"))
    lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{authblk}",
        "",
        f"\\title{{{title}}}",
        "",
        r"\begin{document}",
        r"\maketitle",
        "",
    ]

    order = get_template_sections_order(template_id)
    if not order:
        order = [k for k in sections_data if k != "title"]

    for key in order:
        if key not in sections_data:
            continue
        text = sections_data[key]
        if isinstance(text, dict):
            text = text.get("text", "")
        if not text:
            continue
        sec_title = format_section_title(key)
        lines.append(f"\\section{{{sec_title}}}")
        lines.append(_escape_latex(text))
        lines.append("")

    lines.append(r"\end{document}")

    filename = f"manuscript_{uuid.uuid4().hex[:8]}.tex"
    filepath = os.path.join("/tmp", filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return filepath


def export_to_plain_text(sections_data: dict) -> str:
    title = sections_data.get("title", "Manuscript Draft")
    lines = [title.upper(), "=" * len(title), ""]

    for key, content in sections_data.items():
        if key == "title":
            continue
        text = content
        if isinstance(text, dict):
            text = text.get("text", "")
        if not text:
            continue
        heading = format_section_title(key).upper()
        lines += [heading, "-" * len(heading), text, ""]

    filename = f"manuscript_{uuid.uuid4().hex[:8]}.txt"
    filepath = os.path.join("/tmp", filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return filepath


def create_download_record(
    filepath: str, filename: str, user_id: str, template_id: str, sections: list
) -> dict:
    export_id = "exp_h" + uuid.uuid4().hex[:8]
    now = datetime.utcnow()
    expires_at = now + timedelta(days=7)

    record = {
        "export_id": export_id,
        "user_id": user_id,
        "template_id": template_id,
        "sections_exported": ",".join(sections),
        "filename": filename,
        "filepath": filepath,
        "created_at": now.strftime("%Y-%m-%d %H:%M:%S"),
        "download_count": "0",
        "expires_at": expires_at.strftime("%Y-%m-%d %H:%M:%S"),
    }

    df = _read_history()
    df = pd.concat([df, pd.DataFrame([record])], ignore_index=True)
    _write_history(df)

    return {
        "download_id": export_id,
        "expires_at": expires_at.strftime("%Y-%m-%d %H:%M"),
        "filename": filename,
        "download_url": f"/export/download/{export_id}",
    }


def get_user_export_history(user_id: str) -> list:
    df = _read_history()
    rows = df[df["user_id"] == user_id].to_dict(orient="records")
    now = datetime.utcnow()
    for r in rows:
        try:
            exp = datetime.strptime(r["expires_at"], "%Y-%m-%d %H:%M:%S")
            r["is_expired"] = exp < now
            r["expires_in_days"] = max(0, (exp - now).days)
        except Exception:
            r["is_expired"] = True
            r["expires_in_days"] = 0
    rows.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return rows


def get_export_record(export_id: str) -> dict | None:
    df = _read_history()
    row = df[df["export_id"] == export_id]
    return row.iloc[0].to_dict() if not row.empty else None


def increment_download_count(export_id: str) -> None:
    df = _read_history()
    mask = df["export_id"] == export_id
    if mask.any():
        current = int(df.loc[mask, "download_count"].iloc[0] or 0)
        df.loc[mask, "download_count"] = str(current + 1)
        _write_history(df)


def check_reader_rate_limit(user_id: str) -> bool:
    df = _read_history()
    today = datetime.utcnow().strftime("%Y-%m-%d")
    count = len(df[(df["user_id"] == user_id) & (df["created_at"].str.startswith(today))])
    return count < 5


def cleanup_expired_exports() -> int:
    df = _read_history()
    if df.empty:
        return 0
    now = datetime.utcnow()

    def _is_expired(val: str) -> bool:
        try:
            return datetime.strptime(val, "%Y-%m-%d %H:%M:%S") < now
        except Exception:
            return True

    expired_mask = df["expires_at"].apply(_is_expired)
    for _, row in df[expired_mask].iterrows():
        try:
            if os.path.exists(row["filepath"]):
                os.remove(row["filepath"])
        except Exception as e:
            logger.warning("Could not remove %s: %s", row.get("filepath"), e)

    removed = int(expired_mask.sum())
    _write_history(df[~expired_mask].reset_index(drop=True))
    return removed

# ── Journal templates ─────────────────────────────────────────────────────────

JOURNAL_TEMPLATES_CSV = os.path.join(config.CSV_DIR, "journal_templates.csv")
EXPORT_JOBS_CSV = os.path.join(config.CSV_DIR, "export_jobs.csv")

JOURNAL_TEMPLATES_COLS = [
    "template_id", "journal_name", "publisher", "format", "font_family",
    "font_size", "line_spacing", "margin_top", "margin_bottom",
    "margin_left", "margin_right", "reference_style", "author_format",
    "affiliation_format", "abstract_limit", "keywords_required",
    "competing_interests_format", "funding_format",
    "sections_order", "special_requirements",
]
EXPORT_JOBS_COLS = [
    "job_id", "user_id", "export_type", "status", "template_id",
    "created_at", "completed_at", "file_paths", "error_message",
]

_SEED_JOURNAL_TEMPLATES = [
    {
        "template_id": "jnl_001", "journal_name": "PLoS Pathog", "publisher": "PLOS",
        "format": "single_column", "font_family": "Arial", "font_size": "11",
        "line_spacing": "1.5", "margin_top": "25mm", "margin_bottom": "25mm",
        "margin_left": "20mm", "margin_right": "20mm", "reference_style": "PLOS",
        "author_format": "Last FM", "affiliation_format": "numbered_list",
        "abstract_limit": "300", "keywords_required": "true",
        "competing_interests_format": "Competing Interests",
        "funding_format": "Funding Information",
        "sections_order": "author_order,methods,acknowledgments,funding,competing_interests,credit",
        "special_requirements": "Author summary required after abstract",
    },
    {
        "template_id": "jnl_002", "journal_name": "PNAS", "publisher": "PNAS",
        "format": "single_column", "font_family": "Times New Roman", "font_size": "10",
        "line_spacing": "1.0", "margin_top": "20mm", "margin_bottom": "20mm",
        "margin_left": "20mm", "margin_right": "20mm", "reference_style": "PNAS",
        "author_format": "F.M. Last", "affiliation_format": "superscript_numbers",
        "abstract_limit": "250", "keywords_required": "true",
        "competing_interests_format": "No competing interests section",
        "funding_format": "Author contributions must include specific roles",
        "sections_order": "author_order,methods,acknowledgments,credit,funding",
        "special_requirements": "Significance statement required; max 120 words",
    },
    {
        "template_id": "jnl_003", "journal_name": "Mol Neurobiol", "publisher": "Springer",
        "format": "single_column", "font_family": "Times New Roman", "font_size": "12",
        "line_spacing": "1.5", "margin_top": "25mm", "margin_bottom": "25mm",
        "margin_left": "25mm", "margin_right": "25mm", "reference_style": "Springer",
        "author_format": "Last FM", "affiliation_format": "superscript_numbers",
        "abstract_limit": "300", "keywords_required": "true",
        "competing_interests_format": "Declarations section includes competing interests",
        "funding_format": "Funding section",
        "sections_order": "author_order,methods,acknowledgments,funding,competing_interests,credit",
        "special_requirements": "Springer format with declarations section",
    },
    {
        "template_id": "jnl_004", "journal_name": "Brain Pathol", "publisher": "Wiley",
        "format": "single_column", "font_family": "Times New Roman", "font_size": "12",
        "line_spacing": "2.0", "margin_top": "25mm", "margin_bottom": "25mm",
        "margin_left": "25mm", "margin_right": "25mm", "reference_style": "Wiley",
        "author_format": "Last F.M.", "affiliation_format": "superscript_numbers",
        "abstract_limit": "250", "keywords_required": "true",
        "competing_interests_format": "Conflict of Interest",
        "funding_format": "Grant Information",
        "sections_order": "author_order,methods,acknowledgments,competing_interests,funding,credit",
        "special_requirements": "Medical journal format; structured abstract may be required",
    },
    {
        "template_id": "jnl_005", "journal_name": "Acta Neuropathol Commun", "publisher": "BMC",
        "format": "single_column", "font_family": "Arial", "font_size": "11",
        "line_spacing": "1.5", "margin_top": "25mm", "margin_bottom": "25mm",
        "margin_left": "25mm", "margin_right": "25mm", "reference_style": "BMC",
        "author_format": "Last FM", "affiliation_format": "affiliation_list",
        "abstract_limit": "350", "keywords_required": "true",
        "competing_interests_format": "Competing interests",
        "funding_format": "Funding",
        "sections_order": "author_order,methods,acknowledgments,funding,competing_interests,credit",
        "special_requirements": "BMC open access; background and conclusions sections",
    },
    {
        "template_id": "jnl_006", "journal_name": "Nat Commun", "publisher": "Nature Portfolio",
        "format": "single_column", "font_family": "Arial", "font_size": "11",
        "line_spacing": "1.5", "margin_top": "25mm", "margin_bottom": "25mm",
        "margin_left": "20mm", "margin_right": "20mm", "reference_style": "Nature",
        "author_format": "Last F.M.", "affiliation_format": "superscript_numbers",
        "abstract_limit": "200", "keywords_required": "false",
        "competing_interests_format": "Competing interests",
        "funding_format": "Funding",
        "sections_order": "author_order,methods,acknowledgments,competing_interests,funding,credit",
        "special_requirements": "Nature family format; abstract max 200 words, no subheadings",
    },
]

_GENERIC_JOURNAL_NAMES = {"Nature", "Cell", "PLOS ONE", "Science"}


def _read_journal_templates() -> pd.DataFrame:
    if not os.path.exists(JOURNAL_TEMPLATES_CSV):
        return pd.DataFrame(columns=JOURNAL_TEMPLATES_COLS)
    return pd.read_csv(JOURNAL_TEMPLATES_CSV, dtype=str).fillna("")


def _read_export_jobs() -> pd.DataFrame:
    if not os.path.exists(EXPORT_JOBS_CSV):
        return pd.DataFrame(columns=EXPORT_JOBS_COLS)
    return pd.read_csv(EXPORT_JOBS_CSV, dtype=str).fillna("")


def _write_export_jobs(df: pd.DataFrame) -> None:
    df.to_csv(EXPORT_JOBS_CSV, index=False, quoting=csv.QUOTE_ALL)


def _migrate_to_real_journals() -> None:
    """Replace generic placeholder journals with real lab journal templates."""
    if not os.path.exists(JOURNAL_TEMPLATES_CSV):
        return
    try:
        df = pd.read_csv(JOURNAL_TEMPLATES_CSV, dtype=str).fillna("")
        if df.empty:
            return
        names = set(df["journal_name"].tolist())
        if names & _GENERIC_JOURNAL_NAMES:
            pd.DataFrame(_SEED_JOURNAL_TEMPLATES).to_csv(
                JOURNAL_TEMPLATES_CSV, index=False, quoting=csv.QUOTE_ALL
            )
            logger.info("Journal templates migrated from generic to real lab journals.")
    except Exception as e:
        logger.warning("Journal template migration failed: %s", e)


def _seed_journal_templates_if_empty() -> None:
    if os.path.exists(JOURNAL_TEMPLATES_CSV):
        try:
            df = pd.read_csv(JOURNAL_TEMPLATES_CSV, dtype=str).fillna("")
            if not df.empty:
                _migrate_to_real_journals()
                return
        except Exception:
            pass
    pd.DataFrame(_SEED_JOURNAL_TEMPLATES).to_csv(
        JOURNAL_TEMPLATES_CSV, index=False, quoting=csv.QUOTE_ALL
    )
    logger.info("Journal templates seeded.")


def _seed_export_jobs_if_empty() -> None:
    if not os.path.exists(EXPORT_JOBS_CSV):
        pd.DataFrame(columns=EXPORT_JOBS_COLS).to_csv(
            EXPORT_JOBS_CSV, index=False, quoting=csv.QUOTE_ALL
        )


def get_journal_templates() -> list:
    return _read_journal_templates().to_dict(orient="records")


def get_journal_template(template_id: str) -> dict | None:
    df = _read_journal_templates()
    row = df[df["template_id"] == template_id]
    return row.iloc[0].to_dict() if not row.empty else None


# ── Export jobs ───────────────────────────────────────────────────────────────

def create_export_job(job_id: str, user_id: str, export_type: str, template_id: str) -> None:
    record = {
        "job_id": job_id, "user_id": user_id, "export_type": export_type,
        "status": "running", "template_id": template_id,
        "created_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "completed_at": "", "file_paths": "", "error_message": "",
    }
    df = _read_export_jobs()
    df = pd.concat([df, pd.DataFrame([record])], ignore_index=True)
    _write_export_jobs(df)


def complete_export_job(job_id: str, file_paths: dict) -> None:
    df = _read_export_jobs()
    mask = df["job_id"] == job_id
    if mask.any():
        df.loc[mask, "status"] = "completed"
        df.loc[mask, "completed_at"] = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        df.loc[mask, "file_paths"] = str(file_paths)
        _write_export_jobs(df)


def error_export_job(job_id: str, error_msg: str) -> None:
    df = _read_export_jobs()
    mask = df["job_id"] == job_id
    if mask.any():
        df.loc[mask, "status"] = "error"
        df.loc[mask, "error_message"] = error_msg[:500]
        _write_export_jobs(df)


# ── PDF / HTML generation ─────────────────────────────────────────────────────

import html as _html_lib


def _safe_html(text: str) -> str:
    return _html_lib.escape(str(text or "")).replace("\n", "<br/>")


def get_default_pdf_styles() -> str:
    return """
        @page { size: A4; margin: 25mm; }
        body { font-family: 'Times New Roman', serif; font-size: 12pt;
               line-height: 1.5; color: #000; }
        h1 { font-size: 16pt; font-weight: bold; margin: 20pt 0 10pt; }
        h2 { font-size: 13pt; font-weight: bold; margin: 14pt 0 7pt; }
        p { margin: 8pt 0; text-align: justify; }
        .section-block { margin-bottom: 14pt; }
    """


def get_journal_css(template_id: str) -> str:
    tmpl = get_journal_template(template_id)
    if not tmpl:
        return get_default_pdf_styles()

    font = tmpl.get("font_family", "Times New Roman")
    size = tmpl.get("font_size", "12")
    spacing = tmpl.get("line_spacing", "1.5")
    mt = tmpl.get("margin_top", "25mm")
    mb = tmpl.get("margin_bottom", "25mm")
    ml = tmpl.get("margin_left", "25mm")
    mr = tmpl.get("margin_right", "25mm")

    css = f"""
        @page {{ size: A4; margin: {mt} {mr} {mb} {ml}; }}
        body {{ font-family: '{font}', serif; font-size: {size}pt;
               line-height: {spacing}; color: #000; }}
        h1 {{ font-size: calc({size}pt + 4pt); font-weight: bold;
              margin: 18pt 0 9pt; }}
        h2 {{ font-size: calc({size}pt + 2pt); font-weight: bold;
              margin: 13pt 0 6pt; }}
        p {{ margin: 7pt 0; text-align: justify; }}
        .section-block {{ margin-bottom: 12pt; }}
    """

    jname = tmpl.get("journal_name", "")
    if jname == "PLoS Pathog":
        css += """
        h1 { text-align: left; }
        .competing-block { background: #fffacd; padding: 8pt; margin: 10pt 0;
            border-left: 3pt solid #d4a017; }
        .funding-block { background: #f0fff0; padding: 8pt; margin: 10pt 0;
            border-left: 3pt solid #2e7d32; }
        .author-summary-block { background: #e8f4fd; padding: 8pt; margin: 10pt 0;
            border-left: 3pt solid #1565c0; }
        """
    elif jname == "PNAS":
        css += """
        h1 { text-align: center; font-size: calc({size}pt + 3pt); }
        .author-block { text-align: center; font-size: 9pt; margin: 5pt 0; }
        .significance-block { background: #f5f5f5; padding: 8pt; margin: 12pt 0;
            border: 1pt solid #bbb; font-style: italic; }
        .methods-block { font-size: 9pt; line-height: 1.3; }
        """
    elif jname == "Mol Neurobiol":
        css += """
        .declarations-block { background: #fafafa; padding: 7pt; margin: 9pt 0;
            border-left: 3pt solid #555; }
        .competing-block { background: #fafafa; padding: 7pt; margin: 9pt 0; }
        """
    elif jname == "Brain Pathol":
        css += """
        h1 { text-align: left; }
        .conflict-interest-block { background: #fff8e1; padding: 8pt; margin: 10pt 0;
            border-left: 3pt solid #f9a825; }
        .grant-info-block { background: #e8f5e9; padding: 7pt; margin: 9pt 0; }
        """
    elif jname == "Acta Neuropathol Commun":
        css += """
        .competing-block { background: #fffacd; padding: 7pt; margin: 9pt 0; }
        .funding-block { background: #f0fff0; padding: 7pt; margin: 9pt 0; }
        .background-block { font-weight: bold; margin-bottom: 6pt; }
        """
    elif jname == "Nat Commun":
        css += """
        h1 { text-align: center; }
        .author-block { text-align: center; font-size: 10pt; margin: 6pt 0; }
        .methods-block { font-size: 10pt; }
        .competing-block { background: #f5f5f5; padding: 8pt; margin: 10pt 0;
            border-left: 3pt solid #0077b6; }
        """

    return css


def export_to_html(sections_data: dict, template_id: str = "jnl_001") -> str:
    tmpl = get_journal_template(template_id)
    jname = tmpl["journal_name"] if tmpl else "Generic"
    order = get_template_sections_order(template_id)
    if not order:
        order = [k for k in sections_data if k != "title"]

    parts = [f'<div class="manuscript"><h1 class="doc-title">{_safe_html(sections_data.get("title", "Manuscript Draft"))}</h1>']

    _section_headings = {
        "author_order": "Authors &amp; Affiliations",
        "methods": "Materials and Methods",
        "funding": "Funding",
        "acknowledgments": "Acknowledgments",
        "credit": "Author Contributions",
        "competing_interests": "Competing Interests",
    }
    _journal_competing_headings = {
        "jnl_004": "Conflict of Interest",
        "jnl_003": "Declarations",
    }
    _journal_funding_headings = {
        "jnl_004": "Grant Information",
    }
    if template_id in _journal_competing_headings:
        _section_headings["competing_interests"] = _journal_competing_headings[template_id]
    if template_id in _journal_funding_headings:
        _section_headings["funding"] = _journal_funding_headings[template_id]

    _section_classes = {
        "author_order": "author-block",
        "methods": "methods-block",
        "funding": "funding-block",
        "acknowledgments": "ack-block",
        "credit": "credit-block",
        "competing_interests": "competing-block",
    }
    if template_id == "jnl_004":
        _section_classes["competing_interests"] = "conflict-interest-block"
        _section_classes["funding"] = "grant-info-block"
    elif template_id == "jnl_003":
        _section_classes["competing_interests"] = "declarations-block"

    for key in order:
        if key not in sections_data or not sections_data[key]:
            continue
        text = sections_data[key]
        if isinstance(text, dict):
            text = text.get("text", "")
        if not text:
            continue
        heading = _section_headings.get(key, format_section_title(key))
        cls = _section_classes.get(key, "section-block")
        parts.append(f'<h2>{heading}</h2><div class="{cls} section-block"><p>{_safe_html(text)}</p></div>')

    parts.append("</div>")
    return "\n".join(parts)


def generate_pdf_weasyprint(sections_data: dict, template_id: str = "jnl_001") -> str:
    try:
        import weasyprint
    except ImportError:
        raise RuntimeError("WeasyPrint is not installed. Install it with: pip install weasyprint")

    html_body = export_to_html(sections_data, template_id)
    css = get_journal_css(template_id)
    tmpl = get_journal_template(template_id)
    jname = tmpl["journal_name"] if tmpl else "Manuscript"

    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<title>{_safe_html(sections_data.get('title', 'Manuscript'))}</title>
<style>{css}</style>
</head>
<body>
<div class="journal-label" style="font-size:8pt;color:#666;margin-bottom:16pt">
  Formatted for: <strong>{_safe_html(jname)}</strong>
</div>
{html_body}
</body>
</html>"""

    pdf_bytes = weasyprint.HTML(string=full_html).write_pdf()
    filename = f"manuscript_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join("/tmp", filename)
    with open(filepath, "wb") as f:
        f.write(pdf_bytes)
    return filepath


# ── Journal requirements validation ──────────────────────────────────────────

_JOURNAL_REQUIRED: dict[str, list] = {
    # PLoS Pathog — requires author summary, competing interests, funding
    "jnl_001": ["author_order", "methods", "acknowledgments", "competing_interests", "funding"],
    # PNAS — requires credit (author contributions), significance statement is editorial
    "jnl_002": ["author_order", "methods", "acknowledgments", "credit"],
    # Mol Neurobiol — Springer declarations: competing interests mandatory
    "jnl_003": ["author_order", "methods", "competing_interests"],
    # Brain Pathol — Wiley: author order and methods required
    "jnl_004": ["author_order", "methods", "acknowledgments"],
    # Acta Neuropathol Commun — BMC: competing interests mandatory
    "jnl_005": ["author_order", "methods", "competing_interests"],
    # Nat Commun — Nature Portfolio: competing interests + funding required
    "jnl_006": ["author_order", "methods", "competing_interests", "funding"],
}
_JOURNAL_RECOMMENDED: dict[str, list] = {
    "jnl_001": ["credit"],
    "jnl_002": ["funding", "competing_interests"],
    "jnl_003": ["funding", "credit", "acknowledgments"],
    "jnl_004": ["funding", "competing_interests", "credit"],
    "jnl_005": ["funding", "acknowledgments", "credit"],
    "jnl_006": ["acknowledgments", "credit"],
}

_SECTION_NAMES = {
    "author_order": "Author Order & Affiliations",
    "methods": "Materials and Methods",
    "funding": "Funding",
    "acknowledgments": "Acknowledgments",
    "credit": "Author Contributions (CRediT)",
    "competing_interests": "Competing Interests",
}


def validate_journal_requirements(present_sections: list, template_id: str) -> list:
    tmpl = get_journal_template(template_id)
    if not tmpl:
        return [{"status": "fail", "message": f"Unknown template: {template_id}"}]

    results = []
    jname = tmpl.get("journal_name", template_id)
    results.append({
        "status": "info",
        "message": f"Validating for {jname} ({tmpl.get('publisher', '')})",
    })

    required = _JOURNAL_REQUIRED.get(template_id, [])
    recommended = _JOURNAL_RECOMMENDED.get(template_id, [])

    for sec in required:
        name = _SECTION_NAMES.get(sec, sec)
        if sec in present_sections:
            results.append({"status": "pass", "message": f"{name}: present (required)"})
        else:
            results.append({"status": "fail", "message": f"{name}: MISSING (required for {jname})"})

    for sec in recommended:
        name = _SECTION_NAMES.get(sec, sec)
        if sec in present_sections:
            results.append({"status": "pass", "message": f"{name}: present (recommended)"})
        else:
            results.append({"status": "warning", "message": f"{name}: missing (recommended for {jname})"})

    kw_req = tmpl.get("keywords_required", "false").lower() == "true"
    if kw_req:
        results.append({"status": "warning", "message": f"{jname} requires Keywords — add them manually to the final manuscript."})

    abstract_limit = tmpl.get("abstract_limit", "")
    if abstract_limit:
        results.append({"status": "info", "message": f"Abstract word limit: {abstract_limit} words."})

    if template_id == "jnl_002":
        results.append({"status": "warning", "message": "PNAS requires a Significance Statement (max 120 words) — add before submission."})

    if template_id == "jnl_001":
        results.append({"status": "warning", "message": "PLoS Pathog requires an Author Summary paragraph after the Abstract."})

    special = tmpl.get("special_requirements", "").strip()
    if special and special.lower() not in ("", "null"):
        results.append({"status": "info", "message": f"Special: {special}"})

    return results


# ── Updated bootstrap ─────────────────────────────────────────────────────────

def _original_bootstrap():
    pass  # placeholder reference


_orig_bootstrap = bootstrap_export_schema


def bootstrap_export_schema() -> None:
    _orig_bootstrap()
    _seed_journal_templates_if_empty()
    _seed_export_jobs_if_empty()


# ── APScheduler integration ───────────────────────────────────────────────────

_scheduler = None


def init_scheduler(app) -> None:
    global _scheduler
    if _scheduler is not None:
        return

    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        import atexit

        _scheduler = BackgroundScheduler(daemon=True)

        def _cleanup_job():
            with app.app_context():
                try:
                    removed = cleanup_expired_exports()
                    logger.info("Scheduler: cleaned %d expired exports", removed)
                except Exception as e:
                    logger.error("Scheduler cleanup error: %s", e)

        def _cloud_backup_job():
            logger.info("Scheduler: cloud backup job triggered (Google Drive placeholder)")

        _scheduler.add_job(
            func=_cleanup_job,
            trigger="interval",
            hours=6,
            id="cleanup_exports",
            replace_existing=True,
        )
        _scheduler.add_job(
            func=_cloud_backup_job,
            trigger="interval",
            hours=24,
            id="cloud_backup",
            replace_existing=True,
        )
        _scheduler.start()
        atexit.register(lambda: _scheduler.shutdown(wait=False))
        logger.info("Background scheduler started (cleanup every 6h, backup every 24h)")
    except Exception as e:
        logger.warning("Could not start background scheduler: %s", e)
