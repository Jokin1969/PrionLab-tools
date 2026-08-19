"""Downloadable report of a library-chat conversation — PDF (ReportLab,
same stack as the Journal Club report, see services/jc.py for why not
WeasyPrint) and Word (python-docx, same stack as refs_exporter.py, so
the text is real selectable/editable content, not an image).

Both formats render the same content: every question/answer pair up to
(and including) the point the operator asked for the report, plus the
numbered PrionVault sources cited in each answer.
"""
from __future__ import annotations

import html as _html
import re
from datetime import datetime
from typing import Optional

ACCENT_HEX = "0F3460"


def _fmt_date(iso: Optional[str]) -> str:
    if not iso:
        return ""
    try:
        return datetime.fromisoformat(iso.replace("Z", "+00:00")).strftime("%d/%m/%Y %H:%M")
    except ValueError:
        return iso[:16]


def _strip_markdown(text: str) -> str:
    """Answers come formatted with light Markdown (## headings,
    **bold**, [N] citation markers). Both renderers below handle bold
    and headings themselves; this only normalises stray artefacts that
    don't map to either format's own markup."""
    return (text or "").strip()


def _citation_lines(citations: Optional[list[dict]]) -> list[str]:
    if not citations:
        return []
    out = []
    for c in citations:
        bits = [c.get("authors") or "", str(c.get("year") or ""), c.get("journal") or ""]
        meta = " · ".join(b for b in bits if b)
        ident = f"DOI: {c['doi']}" if c.get("doi") else (f"PMID: {c['pubmed_id']}" if c.get("pubmed_id") else "")
        line = f"[{c.get('n')}] {c.get('title') or '(sin título)'}"
        if meta:
            line += f" — {meta}"
        if ident:
            line += f" — {ident}"
        out.append(line)
    return out


# ── PDF (ReportLab) ─────────────────────────────────────────────────────────

def render_pdf(chat: dict, messages: list[dict]) -> bytes:
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("cr_h1", parent=styles["Title"], fontSize=18,
                         textColor=colors.HexColor("#" + ACCENT_HEX), spaceAfter=2)
    subtitle = ParagraphStyle("cr_subtitle", parent=styles["Normal"], fontSize=9.5,
                               textColor=colors.HexColor("#6b7280"), spaceAfter=18)
    q_style = ParagraphStyle("cr_q", parent=styles["Normal"], fontSize=11, leading=15,
                              textColor=colors.HexColor("#065f46"), spaceBefore=14, spaceAfter=6,
                              backColor=colors.HexColor("#ecfdf5"), borderPadding=8)
    a_style = ParagraphStyle("cr_a", parent=styles["Normal"], fontSize=10.5, leading=15,
                              textColor=colors.HexColor("#111827"), spaceAfter=4)
    a_head = ParagraphStyle("cr_a_head", parent=styles["Heading3"], fontSize=11,
                             textColor=colors.HexColor("#" + ACCENT_HEX), spaceBefore=6, spaceAfter=3)
    cite_style = ParagraphStyle("cr_cite", parent=styles["Normal"], fontSize=8.5, leading=12,
                                 textColor=colors.HexColor("#6b7280"), leftIndent=6)
    meta_style = ParagraphStyle("cr_meta", parent=styles["Normal"], fontSize=8.5,
                                 textColor=colors.HexColor("#9ca3af"), spaceAfter=2)

    def esc(s: str) -> str:
        return _html.escape(s or "")

    def md_inline(text: str) -> str:
        # ReportLab Paragraphs read a small XML-like markup — escape
        # first, then re-apply the two Markdown constructs we care about.
        t = esc(text)
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = t.replace("\n", "<br/>")
        return t

    story = [
        Paragraph("Informe de conversación — Chat general", h1),
        Paragraph(f"PrionVault{(' · ' + esc(chat.get('title'))) if chat.get('title') else ''} · "
                  f"generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle),
    ]

    for m in messages:
        if m["role"] == "user":
            story.append(Paragraph(f"🧑 {md_inline(m['content'])}", q_style))
        else:
            story.append(Paragraph("🤖 Respuesta", a_head))
            if m.get("provider_label"):
                story.append(Paragraph(esc(m["provider_label"]), meta_style))
            for block in _strip_markdown(m["content"]).split("\n\n"):
                if block.strip():
                    story.append(Paragraph(md_inline(block), a_style))
            lines = _citation_lines(m.get("citations"))
            if lines:
                story.append(Spacer(1, 4))
                story.append(Paragraph("Fuentes citadas de PrionVault:", meta_style))
                for line in lines:
                    story.append(Paragraph(esc(line), cite_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#e5e7eb"), spaceBefore=10, spaceAfter=4))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.2*cm, rightMargin=2.2*cm)
    doc.build(story)
    return buf.getvalue()


# ── Word (python-docx) ───────────────────────────────────────────────────────

def _article_link(base_url: str, article_id: Optional[str]) -> str:
    return f"{base_url}/prionvault/?open={article_id}" if base_url and article_id else ""


def _reference_lines(references: Optional[list[dict]]) -> list[dict]:
    """Normalises each reference into {display, pv_url, doi_url, pmid_url}
    for both renderers below."""
    from .article_share import _base_url
    base_url = _base_url()
    out = []
    for r in (references or []):
        bits = [r.get("title") or "(título desconocido)"]
        meta = " · ".join(b for b in [r.get("authors"), str(r.get("year") or "") or None,
                                       r.get("journal")] if b)
        if meta:
            bits.append(meta)
        out.append({
            "display":  " — ".join(bits),
            "pv_url":   _article_link(base_url, r.get("id")) if r.get("in_vault") else "",
            "doi_url":  f"https://doi.org/{r['doi']}" if r.get("doi") else "",
            "pmid_url": f"https://pubmed.ncbi.nlm.nih.gov/{r['pubmed_id']}/" if r.get("pubmed_id") else "",
            "doi":      r.get("doi") or "",
            "pmid":     r.get("pubmed_id") or "",
        })
    return out


def render_article_chat_pdf(article: dict, chat: dict, messages: list[dict],
                            references: Optional[list[dict]] = None) -> bytes:
    """PDF of a per-article chat conversation: article header, then each
    Q&A pair as its own heading/answer pair, then a deduplicated
    "Referencias" section for every DOI/PMID the answers mentioned."""
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("ac_h1", parent=styles["Title"], fontSize=17,
                         textColor=colors.HexColor("#" + ACCENT_HEX), spaceAfter=6)
    meta_style = ParagraphStyle("ac_meta", parent=styles["Normal"], fontSize=9.5, leading=14,
                                 textColor=colors.HexColor("#374151"), spaceAfter=2)
    subtitle = ParagraphStyle("ac_subtitle", parent=styles["Normal"], fontSize=9,
                               textColor=colors.HexColor("#6b7280"), spaceAfter=18)
    q_style = ParagraphStyle("ac_q", parent=styles["Heading3"], fontSize=11.5, leading=15,
                              textColor=colors.HexColor("#065f46"), spaceBefore=16, spaceAfter=6,
                              backColor=colors.HexColor("#ecfdf5"), borderPadding=8)
    a_style = ParagraphStyle("ac_a", parent=styles["Normal"], fontSize=10.5, leading=15,
                              textColor=colors.HexColor("#111827"), spaceAfter=4)
    ref_h_style = ParagraphStyle("ac_ref_h", parent=styles["Heading2"], fontSize=13,
                                  textColor=colors.HexColor("#" + ACCENT_HEX), spaceBefore=20, spaceAfter=8)
    ref_style = ParagraphStyle("ac_ref", parent=styles["Normal"], fontSize=9, leading=13,
                                textColor=colors.HexColor("#374151"), spaceAfter=6)

    def esc(s: str) -> str:
        return _html.escape(s or "")

    def md_inline(text: str) -> str:
        t = esc(text)
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = t.replace("\n", "<br/>")
        return t

    meta_bits = []
    if article.get("authors"): meta_bits.append(esc(article["authors"]))
    line2 = [b for b in [str(article.get("year") or "") or None, article.get("journal")] if b]
    if line2: meta_bits.append(esc(" · ".join(line2)))
    if article.get("doi"):       meta_bits.append(f"DOI: {esc(article['doi'])}")
    if article.get("pubmed_id"): meta_bits.append(f"PMID: {esc(article['pubmed_id'])}")

    story = [
        Paragraph("Informe de conversación — Chat de artículo", h1),
        Paragraph(esc(article.get("title") or "(sin título)"), meta_style),
    ]
    for bit in meta_bits:
        story.append(Paragraph(bit, meta_style))
    story.append(Paragraph(
        f"PrionVault{(' · ' + esc(chat.get('title'))) if chat.get('title') else ''} · "
        f"generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}", subtitle))

    q_num = 0
    for m in messages:
        if m["role"] == "user":
            q_num += 1
            story.append(Paragraph(f"Pregunta {q_num}: {md_inline(m['content'])}", q_style))
        else:
            if m.get("provider_label"):
                story.append(Paragraph(esc(m["provider_label"]), meta_style))
            for block in _strip_markdown(m["content"]).split("\n\n"):
                if block.strip():
                    story.append(Paragraph(md_inline(block), a_style))

    refs = _reference_lines(references)
    if refs:
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#e5e7eb"), spaceBefore=10, spaceAfter=2))
        story.append(Paragraph("Referencias", ref_h_style))
        for i, r in enumerate(refs, 1):
            line = f"[{i}] {esc(r['display'])}"
            links = []
            if r["pv_url"]:  links.append(f'<link href="{esc(r["pv_url"])}" color="#0F3460">Ver en PrionVault</link>')
            if r["doi_url"]: links.append(f'<link href="{esc(r["doi_url"])}" color="#0F3460">DOI: {esc(r["doi"])}</link>')
            if r["pmid_url"]:links.append(f'<link href="{esc(r["pmid_url"])}" color="#0F3460">PMID: {esc(r["pmid"])}</link>')
            if links:
                line += "<br/>" + " &nbsp;·&nbsp; ".join(links)
            story.append(Paragraph(line, ref_style))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.2*cm, rightMargin=2.2*cm)
    doc.build(story)
    return buf.getvalue()


def render_article_chat_docx(article: dict, chat: dict, messages: list[dict],
                             references: Optional[list[dict]] = None) -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    accent = RGBColor(0x0F, 0x34, 0x60)
    green  = RGBColor(0x06, 0x5F, 0x46)
    grey   = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_heading("Informe de conversación — Chat de artículo", level=1)
    title.runs[0].font.color.rgb = accent

    art_title = doc.add_paragraph()
    r = art_title.add_run(article.get("title") or "(sin título)")
    r.bold = True
    r.font.size = Pt(11)

    meta_bits = []
    if article.get("authors"): meta_bits.append(article["authors"])
    line2 = [b for b in [str(article.get("year") or "") or None, article.get("journal")] if b]
    if line2: meta_bits.append(" · ".join(line2))
    if article.get("doi"):       meta_bits.append(f"DOI: {article['doi']}")
    if article.get("pubmed_id"): meta_bits.append(f"PMID: {article['pubmed_id']}")
    for bit in meta_bits:
        p = doc.add_paragraph()
        pr = p.add_run(bit)
        pr.font.size = Pt(9.5)
        pr.font.color.rgb = grey

    meta = doc.add_paragraph()
    meta_line = ["PrionVault"]
    if chat.get("title"):
        meta_line.append(chat["title"])
    meta_line.append(f"generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    mr = meta.add_run(" · ".join(meta_line))
    mr.font.size = Pt(9)
    mr.font.color.rgb = grey

    def add_bold_aware(paragraph, text: str):
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    q_num = 0
    for m in messages:
        if m["role"] == "user":
            q_num += 1
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(14)
            hr = h.add_run(f"Pregunta {q_num}: {m['content'] or ''}")
            hr.bold = True
            hr.font.color.rgb = green
        else:
            if m.get("provider_label"):
                pl = doc.add_paragraph()
                plr = pl.add_run(m["provider_label"])
                plr.italic = True
                plr.font.size = Pt(9)
                plr.font.color.rgb = grey
            for block in _strip_markdown(m["content"]).split("\n\n"):
                if block.strip():
                    p = doc.add_paragraph()
                    add_bold_aware(p, block.strip())

    refs = _reference_lines(references)
    if refs:
        doc.add_paragraph("―" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = doc.add_heading("Referencias", level=2)
        rh.runs[0].font.color.rgb = accent
        for i, r in enumerate(refs, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(r["display"])
            link_bits = []
            if r["pv_url"]:  link_bits.append(f"Ver en PrionVault: {r['pv_url']}")
            if r["doi_url"]: link_bits.append(f"DOI: {r['doi_url']}")
            if r["pmid_url"]:link_bits.append(f"PMID: {r['pmid_url']}")
            if link_bits:
                lp = doc.add_paragraph()
                lp.paragraph_format.left_indent = Cm(0.6)
                lr = lp.add_run("  ·  ".join(link_bits))
                lr.font.size = Pt(9)
                lr.font.color.rgb = grey

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_docx(chat: dict, messages: list[dict]) -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    accent = RGBColor(0x0F, 0x34, 0x60)
    green  = RGBColor(0x06, 0x5F, 0x46)
    grey   = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_heading("Informe de conversación — Chat general", level=1)
    title.runs[0].font.color.rgb = accent

    meta = doc.add_paragraph()
    meta_bits = ["PrionVault"]
    if chat.get("title"):
        meta_bits.append(chat["title"])
    meta_bits.append(f"generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r = meta.add_run(" · ".join(meta_bits))
    r.font.size = Pt(9)
    r.font.color.rgb = grey

    def add_bold_aware(paragraph, text: str):
        # Same **bold** convention the answers already use.
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    for m in messages:
        if m["role"] == "user":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            run = p.add_run("🧑 " + (m["content"] or ""))
            run.font.color.rgb = green
            run.bold = True
        else:
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(6)
            hr = h.add_run("🤖 Respuesta" + (f" — {m['provider_label']}" if m.get("provider_label") else ""))
            hr.bold = True
            hr.font.color.rgb = accent
            for block in _strip_markdown(m["content"]).split("\n\n"):
                if block.strip():
                    p = doc.add_paragraph()
                    add_bold_aware(p, block.strip())
            lines = _citation_lines(m.get("citations"))
            if lines:
                cp = doc.add_paragraph()
                cr = cp.add_run("Fuentes citadas de PrionVault:")
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = grey
                for line in lines:
                    lp = doc.add_paragraph(style="List Bullet")
                    lr = lp.add_run(line)
                    lr.font.size = Pt(9)
                    lr.font.color.rgb = grey
            doc.add_paragraph("―" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
