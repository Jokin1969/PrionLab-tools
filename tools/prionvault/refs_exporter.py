"""Generate a formatted .docx from a list of PrionVault references.

Config schema (all keys optional — defaults applied if missing):
{
  "blocks": [
    {
      "id": "authors"|"title"|"journal"|"year"|"doi"|"pmid"|"author_position",
      "active": true,
      "options": { ...block-specific... }
    }
  ],
  "show_labels": false,     // prefix each field with "Authors:", "Title:", etc.
  "show_type": false,       // show "Type: Article|Review" as first field
  "marked_author": "Joaquín Castilla"
}

Per-block options:
  authors:   mode ("all"|"first_et_al"|"first_last"),
             marked_bold/italic/underline (bool), marked_color ("#rrggbb"|null),
             bold/italic/underline/color for other authors
  title:     bold, italic, underline, color
  journal:   bold, italic, underline, color  (default: bold+italic)
  year:      bold, italic, underline, color  (default: bold+italic)
  doi:       with_link (bool), bold, italic, underline, color
  pmid:      with_link (bool), bold, italic, underline, color
  author_position: bold, italic, underline, color
"""
from __future__ import annotations

import io
import re
import unicodedata
import zipfile as _zf
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from lxml import etree

ACCENT = RGBColor(0x0F, 0x34, 0x60)
DIM    = RGBColor(0x6B, 0x72, 0x80)
DARK   = RGBColor(0x1E, 0x2D, 0x3D)

_DEFAULT_BLOCKS: list[dict] = [
    {"id": "authors",         "active": True,  "options": {}},
    {"id": "title",           "active": True,  "options": {"bold": True, "italic": True, "underline": True}},
    {"id": "journal",         "active": True,  "options": {"bold": True, "italic": True}},
    {"id": "year",            "active": True,  "options": {"bold": True, "italic": True}},
    {"id": "doi",             "active": True,  "options": {"with_link": True}},
    {"id": "pmid",            "active": True,  "options": {"with_link": True}},
    {"id": "author_position", "active": False, "options": {}},
]

_LABELS = {
    "authors":         "Authors",
    "title":           "Title",
    "journal":         "Journal",
    "year":            "Year",
    "doi":             "DOI",
    "pmid":            "PMID",
    "author_position": "Position",
    "type":            "Type",
}


# ── Helpers ───────────────────────────────────────────────────────────────

def _norm(s: str) -> str:
    return unicodedata.normalize('NFD', s).encode('ascii', 'ignore').decode('ascii').lower().strip()


def _find_marked_author_index(authors_list: list[str], marked_author: str) -> int | None:
    if not marked_author:
        return None
    parts = marked_author.strip().split()
    if not parts:
        return None
    last  = parts[-1]
    first = parts[0] if len(parts) >= 2 else ''
    initial   = first[0] if first else ''
    norm_last  = _norm(last)
    norm_first = _norm(first) if first else ''

    for i, a in enumerate(authors_list):
        na = _norm(a)
        # Use word boundary to avoid partial matches (Castillo ≠ Castilla)
        if not re.search(r'\b' + re.escape(norm_last) + r'\b', na):
            continue
        # Last name matched — verify first name or initial when available
        if not first:
            return i
        if norm_first and norm_first in na:
            return i
        if initial and re.search(r'\b' + re.escape(initial.lower()) + r'\b', na):
            return i
    return None


def _parse_color(hex_color: str | None) -> RGBColor | None:
    if not hex_color:
        return None
    h = hex_color.lstrip('#')
    if len(h) == 6:
        try:
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        except ValueError:
            pass
    return None


def _add_run(para, text: str, *, bold=False, italic=False, underline=False,
             color: RGBColor | None = None, size=Pt(11)):
    if not text:
        return
    r = para.add_run(text)
    r.font.name      = 'Calibri'
    r.font.size      = size
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.underline = underline
    if color:
        r.font.color.rgb = color


def _sep_run(para):
    _add_run(para, '  ·  ', color=DIM, size=Pt(10))


def _add_hyperlink(para, text: str, url: str, opts: dict):
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    col_hex = (opts.get('color') or '#0F3460').lstrip('#')
    col_el = OxmlElement('w:color')
    col_el.set(qn('w:val'), col_hex)
    rPr.append(col_el)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')   # 11pt = 22 half-points
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '22')
    rPr.append(sz)
    rPr.append(szCs)

    if opts.get('bold'):
        rPr.append(OxmlElement('w:b'))
    if opts.get('italic'):
        rPr.append(OxmlElement('w:i'))

    r.append(rPr)

    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


def _format_authors(
    authors_list: list[str],
    opts: dict,
    marked_author: str,
    mode: str,
) -> list[tuple]:
    """Return [(text, bold, italic, underline, color|None), ...] per displayed author."""
    if mode == 'first_et_al' and len(authors_list) > 1:
        display = [(0, authors_list[0]), (-1, 'et al.')]
    elif mode == 'first_last' and len(authors_list) > 2:
        display = [(0, authors_list[0]), (-1, '…'), (len(authors_list) - 1, authors_list[-1])]
    else:
        display = list(enumerate(authors_list))

    marked_idx = _find_marked_author_index(authors_list, marked_author) if marked_author else None

    result = []
    for orig_idx, name in display:
        is_marked = marked_idx is not None and orig_idx == marked_idx
        if is_marked:
            b   = opts.get('marked_bold',      False)
            it  = opts.get('marked_italic',     False)
            ul  = opts.get('marked_underline',  False)
            col = _parse_color(opts.get('marked_color'))
        else:
            b   = opts.get('bold',      False)
            it  = opts.get('italic',    False)
            ul  = opts.get('underline', False)
            col = _parse_color(opts.get('color'))
        result.append((name, b, it, ul, col))
    return result


# ── Core renderer ─────────────────────────────────────────────────────────

def _render_ref(para, article: dict, config: dict, number: int) -> None:
    blocks        = config.get('blocks', _DEFAULT_BLOCKS)
    show_labels   = config.get('show_labels', False)
    marked_author = config.get('marked_author', '')

    _add_run(para, f'[{number}] ', bold=True, color=ACCENT, size=Pt(10))

    # Determine article type from source_metadata if available
    sm = article.get('source_metadata') or {}
    pub_type = sm.get('publication_type') or sm.get('type') or ''
    type_str = 'Review' if 'review' in pub_type.lower() else 'Article'

    first = True

    if config.get('show_type', False):
        if show_labels:
            _add_run(para, 'Type: ', color=DIM, size=Pt(9))
        _add_run(para, type_str, bold=True, size=Pt(10))
        first = False

    active_blocks = [b for b in blocks if b.get('active', True)]

    for block in active_blocks:
        bid  = block['id']
        opts = block.get('options', {})

        # Resolve value(s)
        if bid == 'authors':
            raw = (article.get('authors') or '').strip()
            authors_list = [a.strip() for a in raw.split(',') if a.strip()] if raw else []
            if not authors_list:
                continue
        elif bid == 'title':
            val = (article.get('title') or '').strip()
            if not val:
                continue
        elif bid == 'journal':
            val = (article.get('journal') or '').strip()
            if not val:
                continue
        elif bid == 'year':
            yr = article.get('year')
            val = str(yr) if yr else ''
            if not val:
                continue
        elif bid == 'doi':
            val = (article.get('doi') or '').strip()
            if not val:
                continue
        elif bid == 'pmid':
            val = (article.get('pubmed_id') or '').strip()
            if not val:
                continue
        elif bid == 'author_position':
            raw = (article.get('authors') or '').strip()
            authors_list_pos = [a.strip() for a in raw.split(',') if a.strip()] if raw else []
            total = len(authors_list_pos)
            if total == 0:
                continue
            midx = _find_marked_author_index(authors_list_pos, marked_author)
            if midx is None:
                continue
            val = f'{midx + 1}/{total}'
        else:
            continue

        if not first:
            _sep_run(para)
        first = False

        label_text = f'{_LABELS.get(bid, bid)}: ' if show_labels else ''

        if bid == 'authors':
            if label_text:
                _add_run(para, label_text, color=DIM, size=Pt(9))
            mode  = opts.get('mode', 'all')
            parts = _format_authors(authors_list, opts, marked_author, mode)
            for j, (name, b, it, ul, col) in enumerate(parts):
                if j > 0:
                    _add_run(para, ', ', size=Pt(10), color=DARK)
                _add_run(para, name, bold=b, italic=it, underline=ul, color=col, size=Pt(10))

        elif bid in ('doi', 'pmid'):
            with_link = opts.get('with_link', True)
            if bid == 'doi':
                url   = f'https://doi.org/{val}' if not val.startswith('http') else val
                label = f'doi:{val}'
            else:
                url   = f'https://pubmed.ncbi.nlm.nih.gov/{val}/'
                label = f'PMID:{val}'

            if label_text:
                _add_run(para, label_text, color=DIM, size=Pt(9))
            if with_link:
                _add_hyperlink(para, label, url, opts)
            else:
                _add_run(para, label,
                         bold=opts.get('bold', False), italic=opts.get('italic', False),
                         underline=opts.get('underline', False),
                         color=_parse_color(opts.get('color')), size=Pt(10))

        elif bid == 'author_position':
            if label_text:
                _add_run(para, label_text, color=DIM)
            _add_run(para, f'({val})',
                     bold=opts.get('bold', False), italic=opts.get('italic', False),
                     underline=opts.get('underline', False),
                     color=_parse_color(opts.get('color')))

        else:
            if label_text:
                _add_run(para, label_text, color=DIM, size=Pt(9))
            _add_run(para, val,
                     bold=opts.get('bold', False), italic=opts.get('italic', False),
                     underline=opts.get('underline', False),
                     color=_parse_color(opts.get('color')), size=Pt(10))


# ── Document builder ──────────────────────────────────────────────────────

def _patch_app_xml(data: bytes) -> bytes:
    APP_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
    root = etree.fromstring(data)
    for tag, val in [('AppVersion', '16.0000'), ('Application', 'Microsoft Office Word')]:
        el = root.find(f'{{{APP_NS}}}{tag}')
        if el is None:
            el = etree.SubElement(root, f'{{{APP_NS}}}{tag}')
        el.text = val
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def generate_refs_docx(articles: list[dict], config: dict | None = None) -> bytes:
    """Return raw .docx bytes for `articles` formatted per `config`."""
    if config is None:
        config = {}
    if 'blocks' not in config:
        import copy
        config['blocks'] = copy.deepcopy(_DEFAULT_BLOCKS)

    doc = Document()

    # Default font: Calibri 11pt
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    tp = doc.add_paragraph()
    r  = tp.add_run('Lista de referencias')
    r.font.name      = 'Calibri'
    r.font.size      = Pt(18)
    r.font.bold      = True
    r.font.color.rgb = ACCENT

    # Subtitle / metadata
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after  = Pt(22)
    sr = sp.add_run(
        f'{len(articles)} referencia{"s" if len(articles) != 1 else ""}  ·  '
        f'Exportado el {datetime.now().strftime("%d/%m/%Y")}'
    )
    sr.font.name      = 'Calibri'
    sr.font.size      = Pt(10)
    sr.font.color.rgb = DIM
    sr.font.italic    = True

    # One paragraph per reference
    for i, art in enumerate(articles, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(8)
        p.paragraph_format.left_indent       = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        _render_ref(p, art, config, i)

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    # Post-process: update AppVersion so Word 365 handles it as modern document
    in_buf  = io.BytesIO(raw)
    out_buf = io.BytesIO()
    with _zf.ZipFile(in_buf, 'r') as zin, \
         _zf.ZipFile(out_buf, 'w', compression=_zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'docProps/app.xml':
                data = _patch_app_xml(data)
            zout.writestr(item, data)

    return out_buf.getvalue()
